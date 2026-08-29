#!/usr/bin/env python3
"""Materialize and deterministically validate native Legal E2E artifacts."""
from __future__ import annotations

import hashlib
import json
import shutil
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def file_digest(path: Path) -> str:
    return "sha256:" + hashlib.sha256(path.read_bytes()).hexdigest()


def docx_text(path: Path) -> str:
    document = Document(path)
    rows = [p.text for p in document.paragraphs]
    rows.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(rows).strip()


def _font(style: Any, name: str, size: float, *, bold: bool | None = None) -> None:
    style.font.name = name; style.font.size = Pt(size)
    rpr = style._element.get_or_add_rPr(); fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name); fonts.set(qn("w:hAnsi"), name)
    if bold is not None: style.font.bold = bold


def apply_business_brief_styles(document: Document) -> None:
    """Resolve the standard_business_brief preset into explicit Word tokens."""
    for section in document.sections:
        section.page_width = Inches(8.5); section.page_height = Inches(11)
        section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
        section.header_distance = section.footer_distance = Inches(.492)
    normal = document.styles["Normal"]; _font(normal, "Calibri", 11)
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6),
                                      ("Heading 3", 12, 8, 4)):
        style = document.styles[name]; _font(style, "Calibri", size, bold=True)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)


def materialize_docx(content: dict[str, Any], output: Path, *, source: Path | None = None) -> dict[str, Any]:
    """Create a brief or append a real amendment section to a protected-source copy."""
    output.parent.mkdir(parents=True, exist_ok=True)
    source_digest = file_digest(source) if source else None
    if source:
        shutil.copy2(source, output); document = Document(output)
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(str(content.get("title") or "Amendment"), level=1)
    else:
        document = Document(); apply_business_brief_styles(document)
        title = document.add_paragraph()
        title.paragraph_format.space_after = Pt(16)
        run = title.add_run(str(content.get("title") or "Legal Memorandum")); run.bold = True; run.font.size = Pt(23)
    sections = content.get("sections")
    if not isinstance(sections, list) or not sections:
        raise ValueError("native DOCX content requires at least one section")
    for section in sections:
        if not isinstance(section, dict) or not str(section.get("body") or "").strip():
            raise ValueError("every DOCX section requires body text")
        heading = str(section.get("heading") or "Analysis")
        document.add_heading(heading, level=2)
        for block in str(section["body"]).split("\n\n"):
            if block.strip(): document.add_paragraph(block.strip())
    document.save(output)
    return validate_docx(output, expected="edit_existing_doc" if source else "make_new_doc",
                         source=source, source_digest=source_digest)


def validate_docx(path: Path, *, expected: str, source: Path | None = None,
                  source_digest: str | None = None) -> dict[str, Any]:
    exists = path.is_file(); valid_zip = exists and zipfile.is_zipfile(path)
    text = docx_text(path) if valid_zip else ""
    digest = file_digest(path) if exists else None
    changed = None if source is None else digest != (source_digest or file_digest(source))
    source_prefix_preserved = None
    if source is not None and valid_zip:
        source_text = docx_text(source)
        source_prefix_preserved = bool(source_text) and text.startswith(source_text[:min(500, len(source_text))])
    valid = bool(exists and valid_zip and text and
                 (expected != "edit_existing_doc" or changed and source_prefix_preserved))
    return {"expected_output": expected, "artifact_exists": exists, "format_valid": valid_zip,
            "nonempty": bool(text), "actually_modified": changed,
            "basic_structure_preserved": source_prefix_preserved, "artifact_valid": valid,
            "artifact_digest": digest, "source_digest": source_digest,
            "text_length": len(text), "artifact_path": str(path)}


def write_manifest(path: Path, value: dict[str, Any]) -> None:
    path.write_text(json.dumps(value, indent=2, sort_keys=True) + "\n")
